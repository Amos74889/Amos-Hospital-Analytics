import os
import io
import datetime
import pandas as pd
import numpy as np
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
from flask_sqlalchemy import SQLAlchemy
from flask_login import (
    LoginManager, UserMixin, login_user,
    login_required, logout_user, current_user
)
from werkzeug.security import generate_password_hash, check_password_hash
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error
import plotly.graph_objects as go
import json
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv
import warnings
warnings.filterwarnings("ignore")

load_dotenv()

# -------------------------------------------------
# APP CONFIGURATION
# -------------------------------------------------

app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "amos-hospital-secret-2024")

database_url = os.environ.get("DATABASE_URL", "sqlite:///local.db")
if database_url.startswith("postgres://"):
    database_url = database_url.replace("postgres://", "postgresql://")

app.config["SQLALCHEMY_DATABASE_URI"] = database_url
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db = SQLAlchemy(app)

# -------------------------------------------------
# LOGIN MANAGER
# -------------------------------------------------

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"
login_manager.login_message = "Please log in to access the dashboard."

# -------------------------------------------------
# DATABASE MODELS
# -------------------------------------------------

class User(UserMixin, db.Model):
    id         = db.Column(db.Integer, primary_key=True)
    username   = db.Column(db.String(120), unique=True, nullable=False)
    password   = db.Column(db.String(255), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)


class HospitalMetric(db.Model):
    id           = db.Column(db.Integer, primary_key=True)
    date         = db.Column(db.Date, nullable=False)
    metric_name  = db.Column(db.String(120), nullable=False)
    metric_value = db.Column(db.Float, nullable=False)


class UploadHistory(db.Model):
    id          = db.Column(db.Integer, primary_key=True)
    filename    = db.Column(db.String(255))
    uploaded_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    row_count   = db.Column(db.Integer)
    user_id     = db.Column(db.Integer, db.ForeignKey("user.id"))


@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))


# -------------------------------------------------
# AUTO MIGRATION — adds missing columns safely
# -------------------------------------------------

def run_migrations():
    try:
        if "sqlite" in database_url:
            import sqlite3
            db_path = database_url.replace("sqlite:///", "")
            if not os.path.isabs(db_path):
                db_path = os.path.join(os.path.dirname(__file__), "instance", db_path)

            if os.path.exists(db_path):
                raw     = sqlite3.connect(db_path)
                cursor  = raw.execute("PRAGMA table_info(user)")
                columns = [row[1] for row in cursor.fetchall()]
                raw.close()

                if "created_at" not in columns:
                    raw = sqlite3.connect(db_path)
                    raw.execute("ALTER TABLE user ADD COLUMN created_at DATETIME DEFAULT CURRENT_TIMESTAMP")
                    raw.commit()
                    raw.close()
                    print("✅ Migration: added 'created_at' column to user table.")
        else:
            from sqlalchemy import text
            with db.engine.connect() as conn:
                result = conn.execute(text("""
                    SELECT column_name FROM information_schema.columns
                    WHERE table_name='user' AND column_name='created_at'
                """))
                if not result.fetchone():
                    conn.execute(text(
                        'ALTER TABLE "user" ADD COLUMN created_at TIMESTAMP DEFAULT NOW()'
                    ))
                    conn.commit()
                    print("✅ Migration: added 'created_at' column to user table (PostgreSQL).")
    except Exception as e:
        print(f"⚠️ Migration warning (non-fatal): {e}")


# -------------------------------------------------
# AUTH ROUTES
# -------------------------------------------------

@app.route("/register", methods=["GET", "POST"])
def register():
    if current_user.is_authenticated:
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")

        if not username or not password:
            flash("Username and password are required.")
            return redirect(url_for("register"))

        if len(password) < 6:
            flash("Password must be at least 6 characters.")
            return redirect(url_for("register"))

        if User.query.filter_by(username=username).first():
            flash("Username already exists. Please choose another.")
            return redirect(url_for("register"))

        try:
            hashed   = generate_password_hash(password)
            new_user = User(username=username, password=hashed)
            db.session.add(new_user)
            db.session.commit()
            flash("Account created successfully! Please sign in.")
            return redirect(url_for("login"))
        except Exception as e:
            db.session.rollback()
            flash(f"Registration error: {str(e)}")
            return redirect(url_for("register"))

    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if current_user.is_authenticated:
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        user     = User.query.filter_by(username=username).first()

        if user and check_password_hash(user.password, password):
            login_user(user, remember=True)
            return redirect(url_for("dashboard"))
        else:
            flash("Invalid username or password. Please try again.")

    return render_template("login.html")


@app.route("/logout")
@login_required
def logout():
    logout_user()
    flash("You have been logged out.")
    return redirect(url_for("login"))


# -------------------------------------------------
# HELPER: RANDOM FOREST MODEL
# -------------------------------------------------

def run_random_forest(df_grouped, top_disease):
    top_df = df_grouped[df_grouped["metric"] == top_disease].copy()
    top_df["MonthIndex"] = range(1, len(top_df) + 1)

    X = top_df[["MonthIndex"]]
    y = top_df["value"]

    if len(top_df) < 4:
        return {
            "rf_mae": 0, "rf_rmse": 0, "rf_r2": 0,
            "predicted_total": int(y.sum()),
            "predictions": y.values.tolist()
        }

    model = RandomForestRegressor(n_estimators=200, random_state=42, max_depth=5)
    model.fit(X, y)
    preds = model.predict(X)

    mae  = round(mean_absolute_error(y, preds), 2)
    rmse = round(np.sqrt(mean_squared_error(y, preds)), 2)
    r2   = round(max(r2_score(y, preds) * 100, 0), 1)

    future_idx   = pd.DataFrame({"MonthIndex": [len(top_df)+1, len(top_df)+2, len(top_df)+3]})
    future_preds = model.predict(future_idx)

    return {
        "rf_mae":          mae,
        "rf_rmse":         rmse,
        "rf_r2":           r2,
        "predicted_total": int(sum(preds)),
        "predictions":     preds.tolist(),
        "rf_forecast":     [round(float(x), 1) for x in future_preds]
    }


# -------------------------------------------------
# HELPER: ARIMA MODEL
# -------------------------------------------------

def run_arima(df_grouped, top_disease):
    top_df = df_grouped[df_grouped["metric"] == top_disease].copy()
    series = top_df["value"].values

    if len(series) < 6:
        return {
            "arima_mae": 0, "arima_rmse": 0,
            "arima_forecast": [], "arima_status": "Not enough data (need 6+ months)"
        }

    try:
        from statsmodels.tsa.arima.model import ARIMA as ARIMAModel

        train_size = max(int(len(series) * 0.8), 4)
        train, test = series[:train_size], series[train_size:]

        model_fit = ARIMAModel(train, order=(1, 1, 1)).fit()
        forecast  = model_fit.forecast(steps=max(len(test), 1))

        if len(test) > 0:
            mae  = round(mean_absolute_error(test, forecast[:len(test)]), 2)
            rmse = round(np.sqrt(mean_squared_error(test, forecast[:len(test)])), 2)
        else:
            mae, rmse = 0, 0

        future_fit = ARIMAModel(series, order=(1, 1, 1)).fit()
        future_fc  = future_fit.forecast(steps=3)

        return {
            "arima_mae":      mae,
            "arima_rmse":     rmse,
            "arima_forecast": [round(float(x), 1) for x in future_fc.tolist()],
            "arima_status":   "Success"
        }

    except Exception as e:
        return {
            "arima_mae": 0, "arima_rmse": 0,
            "arima_forecast": [], "arima_status": f"Error: {str(e)[:50]}"
        }


# -------------------------------------------------
# HELPER: SURGE DETECTION
# -------------------------------------------------

def detect_surges(df, case_distribution):
    alerts = []
    for item in case_distribution:
        metric    = item["metric"]
        metric_df = df[df["metric"] == metric].sort_values("date")
        if len(metric_df) < 3:
            continue

        values   = metric_df["value"].values
        mean_val = np.mean(values[:-1])
        std_val  = np.std(values[:-1])
        last_val = values[-1]

        if std_val == 0:
            continue

        if last_val > mean_val + 1.5 * std_val:
            pct = round(((last_val - mean_val) / mean_val) * 100, 1)
            alerts.append({
                "type":    "danger",
                "metric":  metric,
                "message": f"⚠️ SURGE ALERT: {metric} is {pct}% above average!",
                "value":   int(last_val),
                "mean":    int(mean_val)
            })
        elif last_val > mean_val + 0.8 * std_val:
            alerts.append({
                "type":    "warning",
                "metric":  metric,
                "message": f"⚡ WARNING: {metric} is trending above normal levels.",
                "value":   int(last_val),
                "mean":    int(mean_val)
            })

    return alerts


# -------------------------------------------------
# HELPER: SMART AI PROMPT
# -------------------------------------------------

def build_smart_prompt(data, is_full_report=False):
    alerts       = data.get("alerts", [])
    arima_fc     = data.get("arima_forecast", [])
    top_disease  = data.get("top_disease", "Unknown")
    rf_r2        = data.get("rf_r2", 0)
    rf_mae       = data.get("rf_mae", 0)
    rf_rmse      = data.get("rf_rmse", 0)
    arima_mae    = data.get("arima_mae", 0)
    arima_rmse   = data.get("arima_rmse", 0)
    total_cases  = data.get("total_cases", 0)
    peak_month   = data.get("peak_month", "Unknown")
    date_range   = data.get("date_range", "")
    distribution = data.get("distribution", [])
    rf_forecast  = data.get("rf_forecast", [])

    dist_text   = "\n".join(f"  - {d['metric']}: {d['total']:,} cases ({d['pct']}%)" for d in distribution)
    alerts_text = "\n".join(f"  - {a['message']}" for a in alerts) if alerts else "  - No active surge alerts"

    has_surge        = any(a["type"] == "danger" for a in alerts)
    has_warning      = any(a["type"] == "warning" for a in alerts)
    forecast_rising  = len(arima_fc) >= 2 and arima_fc[-1] > arima_fc[0]
    forecast_falling = len(arima_fc) >= 2 and arima_fc[-1] < arima_fc[0]
    low_accuracy     = rf_r2 < 60
    high_accuracy    = rf_r2 >= 85

    situation_rules = []

    if has_surge:
        situation_rules.append(f"""
⚠️ SURGE ALERT ACTIVE — URGENT RESPONSE REQUIRED:
- STAFF: Immediately increase on-call staff by 25-30% in {top_disease} affected wards
- BEDS: Activate overflow bed protocols — prepare emergency ward expansion within 24hrs
- SUPPLIES: Issue URGENT procurement for {top_disease} medicines — do not wait for scheduled orders
- TRIAGE: Implement fast-track triage to reduce waiting times by at least 40%
- ESCALATION: Notify hospital CEO, county health department, and Kenya CDC immediately
- REPORTING: Submit an incident report to Ministry of Health within 24 hours""")

    if has_warning:
        situation_rules.append(f"""
⚡ WARNING LEVEL — PRECAUTIONARY MEASURES NEEDED:
- STAFF: Place 2-3 additional nurses per shift on standby for the next 2 weeks
- MONITORING: Switch data collection from monthly to weekly for early detection
- SUPPLIES: Pre-order 30-day buffer stock of critical {top_disease} medicines NOW
- MANAGEMENT: Brief all ward managers on surge preparedness protocols
- REVIEW: Schedule daily data review meetings until situation stabilizes""")

    if forecast_rising:
        fc_vals = " → ".join(str(int(f)) for f in arima_fc)
        situation_rules.append(f"""
📈 ARIMA FORECAST IS RISING ({fc_vals}) — PROACTIVE PLANNING REQUIRED:
- BEDS: Plan for {int(arima_fc[-1] * 1.15):,} admissions next month — start bed allocation NOW
- STAFF: Schedule additional rotations 3 weeks before projected peak
- MEDICINE: Place advance orders 3 weeks before forecast peak of {int(arima_fc[-1]):,} cases
- BUDGET: Request emergency budget allocation from hospital administration
- REVIEW: Set up weekly forecast review meetings with department heads""")

    if forecast_falling:
        fc_vals = " → ".join(str(int(f)) for f in arima_fc)
        situation_rules.append(f"""
📉 ARIMA FORECAST IS DECLINING ({fc_vals}) — POSITIVE TREND:
- STAFF: Maintain current staffing — avoid new permanent hires until trend confirmed over 2 months
- SUPPLIES: Reduce next procurement order by 15-20% to minimize waste and costs
- MAINTENANCE: Use this lower-demand period for equipment servicing and staff training
- INVESTIGATION: Document what interventions contributed to the decline for future reference
- PREVENTION: Continue current prevention measures that appear to be working""")

    if low_accuracy:
        situation_rules.append(f"""
🔴 MODEL ACCURACY IS LOW (R²: {rf_r2}%) — DATA IMPROVEMENTS CRITICAL:
- MINIMUM DATA: Current dataset is too small — need at least 24 months of weekly data
- NEW VARIABLES: Add rainfall, temperature, population density, and vaccination rates
- DATA QUALITY: Assign a dedicated data officer to ensure clean, consistent records
- COLLECTION: Switch from monthly to weekly data collection immediately
- FUTURE MODELS: Once more data is available, implement XGBoost or LSTM neural networks
- VALIDATION: Cross-validate model with data from at least 2 other hospitals""")

    if high_accuracy:
        situation_rules.append(f"""
✅ MODEL ACCURACY IS EXCELLENT (R²: {rf_r2}%) — SCALE THE SYSTEM:
- TRUST: The predictive system is reliable — use forecasts for all resource planning decisions
- EXPAND: Roll out ward-level admission predictions for ICU, Maternity, and Emergency
- SHARE: Present this model to county health officials for adoption across all county hospitals
- PUBLISH: Consider publishing findings in an East African medical journal
- AUTOMATE: Set up automated weekly model retraining as new data comes in""")

    disease_rules = {
        "malaria": """
🦟 MALARIA IS THE TOP DISEASE:
- MEDICINE: Stock Artemether-Lumefantrine (AL) — minimum 500 adult courses and 200 pediatric
- DIAGNOSTICS: Ensure 1,000+ RDT (Rapid Diagnostic Test) kits are available
- STAFF: Refresh all nurses on WHO malaria case management guidelines 2023
- PREVENTION: Coordinate with county government for Long-Lasting Insecticidal Nets (LLINs) distribution
- SEASONAL: Malaria peaks March-May and October-November in Kenya — pre-stock before these months
- REPORTING: Report all confirmed cases to DHIS2 within 24 hours as per MOH requirements""",

        "dengue": """
🦟 DENGUE IS THE TOP DISEASE:
- MEDICINE: Stock paracetamol (NOT aspirin/ibuprofen), IV Normal Saline, and IV Ringer's Lactate
- DIAGNOSTICS: Procure NS1 antigen rapid test kits for early confirmation
- MONITORING: Monitor platelet counts twice daily for all dengue patients
- WARNING SIGNS: Train ALL staff to recognize dengue hemorrhagic fever warning signs immediately
- VECTOR CONTROL: Report to Kenya CDC/county vector control unit for area fumigation
- ISOLATION: Nurse dengue patients under mosquito nets to prevent spreading""",

        "cholera": """
💧 CHOLERA IS THE TOP DISEASE — PUBLIC HEALTH EMERGENCY:
- MEDICINE: URGENTLY stock ORS (Oral Rehydration Salts) — minimum 2,000 sachets
- IV FLUIDS: Stock IV Ringer's Lactate and Normal Saline for severe cases
- ISOLATION: Set up dedicated Cholera Treatment Unit (CTU) AWAY from main wards TODAY
- WATER: Report to county water authority IMMEDIATELY — likely contaminated water source
- PPE: Mandatory PPE (gloves, gowns, masks) for ALL staff — cholera is highly contagious
- REPORTING: MANDATORY notifiable disease — report to MOH within 24 hours or face legal penalties
- COMMUNITY: Coordinate with public health officers for community water treatment""",

        "icu": """
🏥 ICU ADMISSIONS ARE HIGH:
- CAPACITY: Review ICU bed capacity — maximum safe capacity is 85% occupancy
- CONVERSION: Convert High Dependency Unit (HDU) beds to ICU standard if needed
- STAFFING: Maintain strict 1:2 nurse-to-patient ratio in ICU — do not compromise this
- EQUIPMENT: Verify availability of ventilators, cardiac monitors, infusion pumps, defibrillators
- REFERRALS: Activate referral protocol with Kenyatta National Hospital or Aga Khan for overflow
- FAMILIES: Implement strict ICU visitation policy to reduce infection risk""",

        "bed": """
🛏️ BED OCCUPANCY IS HIGH:
- THRESHOLD: If occupancy exceeds 85%, implement emergency bed management protocol
- DISCHARGE: Accelerate discharge planning for stable patients — target same-day discharge where safe
- OVERFLOW: Identify overflow areas — consider converting outpatient spaces for inpatient use
- WARDS: Open additional general wards if occupancy exceeds 90% for more than 48 hours
- REFERRALS: Divert non-emergency admissions to nearby facilities
- REPORTING: Report critical bed shortages to county health management team""",
    }

    disease_advice = ""
    for key, advice in disease_rules.items():
        if key.lower() in top_disease.lower():
            disease_advice = advice
            break

    if not disease_advice:
        disease_advice = f"""
🔬 {top_disease} IS THE TOP DISEASE:
- MEDICINE: Conduct immediate audit of {top_disease} treatment medicine stock levels
- PROTOCOLS: Review and distribute current {top_disease} clinical management guidelines to all staff
- MONITORING: Increase monitoring frequency — daily rounds for all {top_disease} patients
- TRAINING: Schedule a {top_disease} case management refresher for nursing and clinical staff
- REPORTING: Ensure all {top_disease} cases are correctly coded in the hospital information system
- SURVEILLANCE: Report unusual patterns to the county disease surveillance officer"""

    situation_block = "\n".join(situation_rules) if situation_rules else "  - All metrics within normal range — maintain current protocols"
    rf_fc_text      = f"RF 3-Month Forecast: {rf_forecast}" if rf_forecast else ""

    sections = f"""1. **Executive Summary** (2-3 sentences of the most critical finding)
2. **Key Findings** (top 3 data-driven insights)
3. **Disease Burden Analysis** (deep focus on {top_disease})
4. **Risk Alerts & Surge Detection** (status of all alerts)
5. **Staff & Resource Allocation** (specific numbers of staff needed)
6. **Medicine & Supply Orders** (name specific medicines with quantities)
7. **Patient Management Advice** (clinical workflow improvements)
8. **Model Performance Analysis** (compare Random Forest vs ARIMA — which is better and why)
9. **Data Collection Improvements** (what data to add for better accuracy)
10. **3-Month Forecast Outlook** (what to expect and how to prepare)""" if is_full_report else """1. **Key Findings**
2. **Active Alerts & Risk Status**
3. **Staff & Resource Allocation** (specific numbers)
4. **Medicine & Supply Orders** (specific medicines)
5. **Patient Management Advice**
6. **Model Comparison** (RF vs ARIMA — which performed better)
7. **Data Improvements**
8. **3-Month Forecast**"""

    prompt = f"""You are Dr. Amara Osei, a senior medical data analyst and hospital resource planning expert
with 15 years of experience in Kenya's healthcare system, working with the Ministry of Health.

=== CURRENT HOSPITAL DATA ===
- Analysis period: {date_range}
- Total reported cases: {total_cases:,}
- Highest disease burden: {top_disease}
- Peak admission month: {peak_month}

=== PREDICTIVE MODEL RESULTS ===
- Random Forest — MAE: {rf_mae}, RMSE: {rf_rmse}, R²: {rf_r2}%
- ARIMA (Time-Series) — MAE: {arima_mae}, RMSE: {arima_rmse}
- ARIMA 3-Month Forecast: {arima_fc}
{rf_fc_text}

=== ACTIVE ALERTS ===
{alerts_text}

=== DISEASE DISTRIBUTION ===
{dist_text}

=== SITUATIONAL ASSESSMENT — FOLLOW THESE RULES STRICTLY ===
{situation_block}

=== DISEASE-SPECIFIC CLINICAL GUIDANCE ===
{disease_advice}

=== YOUR REPORT ===
Write a professional, urgent, and actionable hospital analytics report structured as:
{sections}

STRICT RULES:
1. Use EXACT numbers from the data — never say "many" or "some"
2. Name SPECIFIC medicines with dosages and quantities
3. Give SPECIFIC staff numbers (e.g., "add 4 nurses to the night shift")
4. Reference Kenya-specific context: DHIS2, Kenya CDC, MOH, county health departments
5. If surge detected → use URGENT language and bold key actions
6. If forecast is rising → give CLEAR WARNING with exact preparation timeline
7. Always compare Random Forest vs ARIMA — explain which model to trust and why
8. End with a clear 3-month outlook timeline
9. Be direct, professional, and data-driven — no filler sentences"""

    return prompt


# -------------------------------------------------
# HELPER: CALL GROQ AI  ← replaces Gemini
# -------------------------------------------------

def call_groq(prompt):
    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        raise ValueError("GROQ_API_KEY not set in environment variables")

    client = Groq(api_key=api_key)

    completion = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {
                "role": "system",
                "content": "You are Dr. Amara Osei, a senior medical data analyst with 15 years experience in Kenya's healthcare system."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.7,
        max_tokens=4096,
    )

    return completion.choices[0].message.content


# -------------------------------------------------
# HELPER: BUILD ANALYSIS SUMMARY
# -------------------------------------------------

def build_analysis_summary():
    records = HospitalMetric.query.all()
    if not records:
        return None

    df = pd.DataFrame([{
        "date": r.date, "metric": r.metric_name, "value": r.metric_value
    } for r in records])

    total_cases   = int(df["value"].sum())
    metric_totals = df.groupby("metric")["value"].sum().sort_values(ascending=False)
    top_disease   = metric_totals.index[0]

    df["month"] = pd.to_datetime(df["date"]).dt.strftime("%B")
    peak_month  = df.groupby("month")["value"].sum().idxmax()
    df_grouped  = df.groupby(["date", "metric"])["value"].sum().reset_index()

    rf_results    = run_random_forest(df_grouped, top_disease)
    arima_results = run_arima(df_grouped, top_disease)

    case_distribution = [
        {"metric": m, "total": int(t), "pct": round(t / metric_totals.sum() * 100, 1)}
        for m, t in metric_totals.items()
    ]

    min_date   = df["date"].min()
    max_date   = df["date"].max()
    date_range = f"{pd.to_datetime(min_date).strftime('%b %Y')} – {pd.to_datetime(max_date).strftime('%b %Y')}"
    alerts     = detect_surges(df, case_distribution)

    return {
        "total_cases":       total_cases,
        "top_disease":       top_disease,
        "peak_month":        peak_month,
        "case_distribution": case_distribution,
        "date_range":        date_range,
        "alerts":            alerts,
        "rf_mae":            rf_results["rf_mae"],
        "rf_rmse":           rf_results["rf_rmse"],
        "rf_r2":             rf_results["rf_r2"],
        "predicted_total":   rf_results["predicted_total"],
        "rf_forecast":       rf_results.get("rf_forecast", []),
        "arima_mae":         arima_results["arima_mae"],
        "arima_rmse":        arima_results["arima_rmse"],
        "arima_forecast":    arima_results["arima_forecast"],
        "arima_status":      arima_results["arima_status"],
    }


# -------------------------------------------------
# HELPER: BUILD WORD DOCUMENT
# -------------------------------------------------

def build_word_report(summary, ai_text):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    style           = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title_p           = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run                = title_p.add_run("AMOS HOSPITAL ANALYTICS")
    run.bold           = True
    run.font.size      = Pt(22)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    sub_p               = doc.add_paragraph()
    sub_p.alignment     = WD_ALIGN_PARAGRAPH.CENTER
    run2                = sub_p.add_run("Disease Surveillance & Predictive Analytics Report")
    run2.font.size      = Pt(14)
    run2.font.color.rgb = RGBColor(0x37, 0x47, 0x5A)

    doc.add_paragraph()

    date_p           = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run(f"Report Period: {summary['date_range']}").bold = True

    gen_p           = doc.add_paragraph()
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_p.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y at %H:%M')} | Powered by Groq AI (Llama 3.3 70B)")

    doc.add_paragraph()

    h = doc.add_heading("Summary Statistics", level=1)
    h.runs[0].font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    table       = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers     = ["Total Cases", "Highest Disease", "Peak Month", "Predicted Cases"]
    values      = [
        f"{summary['total_cases']:,}",
        summary["top_disease"],
        summary["peak_month"],
        f"{summary['predicted_total']:,}"
    ]

    for i, (hdr, val) in enumerate(zip(headers, values)):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = cell.paragraphs[0].add_run(hdr)
        hr.bold = True
        hr.font.size = Pt(9)
        hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc  = cell._tc
        tcp = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1565C0")
        shd.set(qn("w:val"), "clear")
        tcp.append(shd)

        vc = table.rows[1].cells[i]
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vc.paragraphs[0].add_run(val)
        vr.bold = True
        vr.font.size = Pt(13)

    doc.add_paragraph()

    h2 = doc.add_heading("Model Performance Metrics", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    mtable       = doc.add_table(rows=1, cols=5)
    mtable.style = "Table Grid"
    mhdrs        = ["Model", "MAE", "RMSE", "R²", "Status"]
    mvals        = [
        ("Random Forest", str(summary["rf_mae"]), str(summary["rf_rmse"]), f"{summary['rf_r2']}%", "✓ Active"),
        ("ARIMA", str(summary["arima_mae"]), str(summary["arima_rmse"]), "N/A", summary["arima_status"]),
    ]

    for i, mh in enumerate(mhdrs):
        cell = mtable.rows[0].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].add_run(mh)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc  = cell._tc
        tcp = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1565C0")
        shd.set(qn("w:val"), "clear")
        tcp.append(shd)

    for row_data in mvals:
        row = mtable.add_row().cells
        for ci, val in enumerate(row_data):
            row[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[ci].paragraphs[0].add_run(val).font.size = Pt(10)

    doc.add_paragraph()

    if summary.get("arima_forecast"):
        h_fc = doc.add_heading("3-Month ARIMA Forecast", level=1)
        h_fc.runs[0].font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
        fc_table       = doc.add_table(rows=2, cols=3)
        fc_table.style = "Table Grid"
        for i, label in enumerate(["Month 1", "Month 2", "Month 3"]):
            fc_table.rows[0].cells[i].paragraphs[0].add_run(label).bold = True
            val = summary["arima_forecast"][i] if i < len(summary["arima_forecast"]) else "N/A"
            fc_table.rows[1].cells[i].paragraphs[0].add_run(
                str(int(val)) if isinstance(val, float) else str(val)
            )

    doc.add_paragraph()

    if summary["alerts"]:
        h3 = doc.add_heading("⚠️ Active Surge Alerts", level=1)
        h3.runs[0].font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        for alert in summary["alerts"]:
            p = doc.add_paragraph()
            r = p.add_run(f"• {alert['message']} (Current: {alert['value']:,} | Avg: {alert['mean']:,})")
            r.bold = True
            r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28) if alert["type"] == "danger" else RGBColor(0xE6, 0x5C, 0x00)

    doc.add_page_break()

    h4 = doc.add_heading("AI-Powered Analysis & Recommendations", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    sub = doc.add_paragraph()
    sub.add_run("Generated by Groq AI (Llama 3.3 70B) · For administrative use only").italic = True
    sub.runs[0].font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E)
    sub.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    for line in ai_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        is_heading = (
            (len(line) > 2 and line[0].isdigit() and ". " in line[:5]) or
            line.startswith("##") or
            (line.startswith("**") and line.endswith("**"))
        )
        if is_heading:
            clean = line.lstrip("#").lstrip("0123456789.").strip().strip("*").strip()
            h = doc.add_heading(clean, level=2)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
        elif line.startswith("- ") or line.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line.lstrip("- •")).font.size = Pt(11)
        else:
            p = doc.add_paragraph()
            p.add_run(line).font.size = Pt(11)

    doc.add_paragraph()
    footer_p           = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr                 = footer_p.add_run(
        f"CONFIDENTIAL — Amos Hospital Analytics · Generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    fr.font.size      = Pt(9)
    fr.font.color.rgb = RGBColor(0x9E, 0x9E, 0x9E)
    fr.italic         = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# -------------------------------------------------
# GENERATE REPORT ROUTE
# -------------------------------------------------

@app.route("/generate-report")
@login_required
def generate_report():
    try:
        summary = build_analysis_summary()
        if not summary:
            flash("No data found. Please upload a CSV first.")
            return redirect(url_for("dashboard"))

        prompt  = build_smart_prompt({
            **summary,
            "distribution": summary["case_distribution"]
        }, is_full_report=True)

        ai_text  = call_groq(prompt)
        doc_buf  = build_word_report(summary, ai_text)
        filename = f"Amos_Hospital_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"

        return send_file(
            doc_buf,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        flash(f"Report error: {str(e)}")
        return redirect(url_for("dashboard"))


# -------------------------------------------------
# AI INSIGHTS ROUTE
# -------------------------------------------------

@app.route("/ai-insights", methods=["POST"])
@login_required
def ai_insights():
    try:
        data = request.get_json()

        smart_data = {
            "alerts":         data.get("alerts", []),
            "arima_forecast": data.get("arima_forecast", []),
            "rf_forecast":    data.get("rf_forecast", []),
            "top_disease":    data.get("top_disease", "Unknown"),
            "rf_r2":          data.get("rf_r2", 0),
            "rf_mae":         data.get("rf_mae", 0),
            "rf_rmse":        data.get("rf_rmse", 0),
            "arima_mae":      data.get("arima_mae", 0),
            "arima_rmse":     data.get("arima_rmse", 0),
            "total_cases":    data.get("total_cases", 0),
            "peak_month":     data.get("peak_month", "Unknown"),
            "date_range":     data.get("date_range", ""),
            "distribution":   data.get("distribution", []),
        }

        prompt  = build_smart_prompt(smart_data, is_full_report=False)
        insight = call_groq(prompt)

        return jsonify({"success": True, "insights": insight})

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# -------------------------------------------------
# HEALTH CHECK ROUTE
# -------------------------------------------------

@app.route("/health")
def health():
    return jsonify({"status": "ok", "time": str(datetime.datetime.utcnow())}), 200


# -------------------------------------------------
# STATS API
# -------------------------------------------------

@app.route("/api/stats")
@login_required
def api_stats():
    try:
        summary = build_analysis_summary()
        if not summary:
            return jsonify({"status": "no_data"}), 200
        return jsonify({
            "status":         "ok",
            "total_cases":    summary["total_cases"],
            "top_disease":    summary["top_disease"],
            "peak_month":     summary["peak_month"],
            "rf_r2":          summary["rf_r2"],
            "arima_forecast": summary["arima_forecast"],
            "alerts_count":   len(summary["alerts"]),
            "date_range":     summary["date_range"],
        }), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# -------------------------------------------------
# CHANGE PASSWORD ROUTE
# -------------------------------------------------

@app.route("/change-password", methods=["POST"])
@login_required
def change_password():
    old_password = request.form.get("old_password", "")
    new_password = request.form.get("new_password", "")
    confirm      = request.form.get("confirm_password", "")

    if not check_password_hash(current_user.password, old_password):
        flash("❌ Current password is incorrect.")
        return redirect(url_for("dashboard"))

    if len(new_password) < 6:
        flash("❌ New password must be at least 6 characters.")
        return redirect(url_for("dashboard"))

    if new_password != confirm:
        flash("❌ New passwords do not match.")
        return redirect(url_for("dashboard"))

    try:
        current_user.password = generate_password_hash(new_password)
        db.session.commit()
        flash("✅ Password changed successfully!")
    except Exception as e:
        db.session.rollback()
        flash(f"❌ Error changing password: {str(e)}")

    return redirect(url_for("dashboard"))


# -------------------------------------------------
# UPLOAD HISTORY ROUTE
# -------------------------------------------------

@app.route("/api/upload-history")
@login_required
def upload_history():
    try:
        history = UploadHistory.query.filter_by(user_id=current_user.id)\
                    .order_by(UploadHistory.uploaded_at.desc()).limit(10).all()
        return jsonify({
            "status": "ok",
            "history": [{
                "filename":    h.filename,
                "uploaded_at": h.uploaded_at.strftime("%Y-%m-%d %H:%M"),
                "row_count":   h.row_count
            } for h in history]
        }), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# -------------------------------------------------
# CLEAR DATA ROUTE
# -------------------------------------------------

@app.route("/clear-data", methods=["POST"])
@login_required
def clear_data():
    try:
        HospitalMetric.query.delete()
        db.session.commit()
        flash("✅ All data cleared successfully. Upload a new CSV to start fresh.")
    except Exception as e:
        db.session.rollback()
        flash(f"❌ Error clearing data: {str(e)}")
    return redirect(url_for("dashboard"))


# -------------------------------------------------
# DASHBOARD
# -------------------------------------------------

@app.route("/", methods=["GET", "POST"])
@login_required
def dashboard():

    if request.method == "POST":
        file = request.files.get("file")
        if not file or file.filename == "":
            flash("No file selected.")
            return redirect(url_for("dashboard"))

        try:
            df = pd.read_csv(file)
            df.columns = df.columns.str.strip()

            date_col = next((c for c in df.columns if "date" in c.lower()), None)
            if not date_col:
                flash("No date column found in CSV. Make sure your file has a 'Date' column.")
                return redirect(url_for("dashboard"))

            df[date_col] = pd.to_datetime(df[date_col], errors="coerce")
            df = df.dropna(subset=[date_col])

            if df.empty:
                flash("No valid dates found in the CSV file.")
                return redirect(url_for("dashboard"))

            HospitalMetric.query.delete()
            rows = 0

            for _, row in df.iterrows():
                for col in df.columns:
                    if col == date_col:
                        continue
                    if pd.api.types.is_numeric_dtype(df[col]) and not pd.isna(row[col]):
                        db.session.add(HospitalMetric(
                            date=row[date_col].date(),
                            metric_name=col,
                            metric_value=float(row[col])
                        ))
                        rows += 1

            history = UploadHistory(
                filename=file.filename,
                row_count=len(df),
                user_id=current_user.id
            )
            db.session.add(history)
            db.session.commit()

            flash(f"✅ Data uploaded successfully! {len(df)} rows, {rows} data points analyzed.")

        except Exception as e:
            db.session.rollback()
            flash(f"Upload error: {str(e)}")

        return redirect(url_for("dashboard"))

    # ── BUILD DASHBOARD ──
    records = HospitalMetric.query.all()
    if not records:
        return render_template("dashboard.html",
            message="No data uploaded yet.",
            username=current_user.username,
            alerts=[], case_distribution=[],
            arima_forecast=[], rf_forecast=[],
            total_cases=0, top_disease="", peak_month="",
            peak_month_pct=0, rf_mae=0, rf_rmse=0, rf_r2=0,
            arima_mae=0, arima_rmse=0, arima_status="",
            predicted_total=0, trend_chart=None, date_range=""
        )

    df = pd.DataFrame([{
        "date": r.date, "metric": r.metric_name, "value": r.metric_value
    } for r in records])

    total_cases   = int(df["value"].sum())
    metric_totals = df.groupby("metric")["value"].sum().sort_values(ascending=False)
    top_disease   = metric_totals.index[0]

    df["month"]    = pd.to_datetime(df["date"]).dt.strftime("%B")
    monthly_totals = df.groupby("month")["value"].sum()
    peak_month     = monthly_totals.idxmax()
    pct_changes    = monthly_totals.pct_change().dropna()
    peak_month_pct = round(pct_changes.max() * 100, 1) if not pct_changes.empty else 0

    df_grouped = df.groupby(["date", "metric"])["value"].sum().reset_index()

    rf_results    = run_random_forest(df_grouped, top_disease)
    arima_results = run_arima(df_grouped, top_disease)

    case_distribution = [
        {"metric": m, "total": int(t), "pct": round(t / metric_totals.sum() * 100, 1)}
        for m, t in metric_totals.items()
    ]

    alerts = detect_surges(df, case_distribution)

    colors    = ["#3B82F6", "#F97316", "#10B981", "#EF4444", "#8B5CF6"]
    trend_fig = go.Figure()
    for i, metric in enumerate(df_grouped["metric"].unique()):
        mdf = df_grouped[df_grouped["metric"] == metric].sort_values("date")
        trend_fig.add_trace(go.Scatter(
            x=mdf["date"].astype(str), y=mdf["value"],
            mode="lines+markers", name=metric,
            line=dict(color=colors[i % len(colors)], width=2.5),
            marker=dict(size=5)
        ))

    trend_fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#CBD5E1", family="DM Sans"),
        legend=dict(orientation="h", y=1.12, x=0.5, xanchor="center"),
        margin=dict(l=40, r=20, t=30, b=40),
        xaxis=dict(gridcolor="rgba(255,255,255,0.05)", showline=False),
        yaxis=dict(gridcolor="rgba(255,255,255,0.05)", showline=False),
        hovermode="x unified"
    )
    trend_chart = json.dumps(trend_fig, default=str)

    min_date   = df["date"].min()
    max_date   = df["date"].max()
    date_range = f"{pd.to_datetime(min_date).strftime('%b %Y')} – {pd.to_datetime(max_date).strftime('%b %Y')}"

    return render_template(
        "dashboard.html",
        username=current_user.username,
        total_cases=f"{total_cases:,}",
        top_disease=top_disease,
        peak_month=peak_month,
        peak_month_pct=peak_month_pct,
        rf_mae=rf_results["rf_mae"],
        rf_rmse=rf_results["rf_rmse"],
        rf_r2=rf_results["rf_r2"],
        predicted_total=f"{rf_results['predicted_total']:,}",
        rf_forecast=rf_results.get("rf_forecast", []),
        arima_mae=arima_results["arima_mae"],
        arima_rmse=arima_results["arima_rmse"],
        arima_forecast=arima_results["arima_forecast"],
        arima_status=arima_results["arima_status"],
        case_distribution=case_distribution,
        alerts=alerts,
        trend_chart=trend_chart,
        date_range=date_range,
        message=None
    )


# -------------------------------------------------
# INITIALIZE DATABASE + RUN MIGRATIONS
# -------------------------------------------------

with app.app_context():
    db.create_all()
    run_migrations()

# -------------------------------------------------
# RUN
# -------------------------------------------------

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)